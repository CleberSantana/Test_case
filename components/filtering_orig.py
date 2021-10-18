import docx
from docx import Document
import os

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

if __name__ == "__main__":
    for file in os.listdir(os.getcwd()):
        if file.endswith(".docm"):
        #print(os.path.join(os.getcwd(), file))
            document = Document(file)        
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for index, paragraph in enumerate(cell.paragraphs):
                            if index == 0:                        
                                if ':]' in paragraph.text:
                                    k=0
                                    for c in paragraph.text:
                                        if ']' in c:
                                            k=k+1
                                    if k == 1:                            
                                        if ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) or 'R1L' in paragraph.text or 'LTM' in paragraph.text or 'LATAM' in paragraph.text:
                                            pass
                                            #print('1- '+paragraph.text)                                    
                                        else:
                                            #print('1- delete - '+paragraph.text)
                                            remove_row(table, row)
                                    if k == 2:                            
                                        if ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) and 'R1L' in paragraph.text or ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) and 'LTM' in paragraph.text or ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) and 'LATAM' in paragraph.text or 'R1L' in paragraph.text and 'LTM' in paragraph.text or 'R1L' in paragraph.text and 'LATAM' in paragraph.text or 'LTM' in paragraph.text and 'LATAM' in paragraph.text:
                                            pass
                                            #print('2 - '+paragraph.text)
                                        else:
                                            #print('2- delete - '+paragraph.text)
                                            remove_row(table, row)
                                    if k == 3:                            
                                        if ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) and 'R1L' in paragraph.text and 'LTM' in paragraph.text or ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) and 'R1L' in paragraph.text and 'LATAM' in paragraph.text or ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) and 'LTM' in paragraph.text and 'LATAM' in paragraph.text or 'R1L' in paragraph.text and 'LTM' in paragraph.text  and 'LATAM' in paragraph.text:
                                            pass
                                            #print('3 - '+paragraph.text)
                                        else:
                                            #print('3- delete - '+paragraph.text)
                                            remove_row(table, row)
                                    if k == 4:                            
                                        if ('allSys: CTS1_2' in paragraph.text or 'AtlMi' in paragraph.text) and 'R1L' in paragraph.text and 'LTM' in paragraph.text and 'LATAM' in paragraph.text :
                                            pass
                                            #print('4 - '+paragraph.text)
                                        else:
                                            #print('4- delete - '+paragraph.text)
                                            remove_row(table, row)
            document.save('FILTRADO_'+file)